from flask import Flask, render_template ,request, send_file, jsonify
from docx import Document
from docx.shared import Inches
from io import BytesIO
import datetime
import uuid
import json
import os


app = Flask(__name__)

# Create directory for saved forms
SAVED_FORMS_DIR = 'saved_forms'
if not os.path.exists(SAVED_FORMS_DIR):
    os.makedirs(SAVED_FORMS_DIR)

# Helper Functions
def is_empty_value(value):
    """Check if a value is empty or contains only whitespace"""
    if value is None:
        return True
    if isinstance(value, str):
        return not value.strip()
    return False

def has_table_data(data, table_id):
    """Check if a table has any non-empty data"""
    i = 0
    while True:
        row_key_base = f"{table_id}_{i}"
        # Check if the first cell of a potential row exists in the form data
        if f"{row_key_base}_0" not in data:
            break
        
        # Check if any cell in the current row has data
        j = 0
        row_has_data = False
        while f"{row_key_base}_{j}" in data:
            if not is_empty_value(data.get(f"{row_key_base}_{j}")):
                row_has_data = True
                break
            j += 1
        
        if row_has_data:
            return True
        i += 1
    return False

def generate_markdown(data, files):
    """Generates a Markdown string from the form data, showing all sections."""
    output = BytesIO()
    writer = lambda s: output.write((s + '\n').encode('utf-8'))

    writer(f"# Incident Report: {data.get('incident_id', 'N/A')}")
    
    # --- Section 1: Triage ---
    writer("\n---\n## 1. Triage & Initial Assessment")
    triage_fields = ['incident_id', 'reported_by', 'date_reported', 'priority', 'summary']
    triage_has_data = any(not is_empty_value(data.get(field)) for field in triage_fields)
    if triage_has_data:
        if not is_empty_value(data.get('incident_id')):
            writer(f"* **Incident ID:** {data['incident_id']}")
        if not is_empty_value(data.get('reported_by')):
            writer(f"* **Reported By:** {data['reported_by']}")
        if not is_empty_value(data.get('date_reported')):
            writer(f"* **Date/Time Reported:** {data['date_reported'].replace('T', ' ')} UTC")
        if not is_empty_value(data.get('priority')):
            writer(f"* **Priority Level:** {data['priority']}")
        if not is_empty_value(data.get('summary')):
            writer(f"* **Initial Summary:**\n    > {data['summary']}")
    else:
        writer("\nNot applicable here.")

    # --- Helper to write tables ---
    def write_table(table_id, headers):
        writer("| " + " | ".join(headers) + " |")
        writer("|" + " :--- |" * len(headers))
        i = 0
        while True:
            row_key_base = f"{table_id}_{i}"
            if f"{row_key_base}_0" not in data: break
            row_data = [data.get(f"{row_key_base}_{j}", "").replace('T', ' ') for j in range(len(headers))]
            if any(cell.strip() for cell in row_data):
                writer("| " + " | ".join(row_data) + " |")
            i += 1

    # --- Section 2: Identification ---
    writer("\n---\n## 2. Identification & Scoping")
    id_table_has_data = has_table_data(data, 'identification_table')
    scope_has_data = not is_empty_value(data.get('scope_summary'))
    iocs_has_data = not is_empty_value(data.get('iocs'))
    if id_table_has_data or scope_has_data or iocs_has_data:
        if id_table_has_data:
            write_table("identification_table", ["Date/Time (UTC)", "Host/Asset", "Source IP", "Destination IP", "Event/Observation", "Analyst Notes"])
        if scope_has_data:
            writer(f"\n**Summary of Scope:**\n{data['scope_summary']}")
        if iocs_has_data:
            writer(f"\n**Indicators of Compromise (IoCs):**\n```\n{data['iocs']}\n```")
        if files and files[0].filename:
            imgstring = ', '.join([f.filename for f in files if f.filename])
            writer(f"\n**Note:** Image evidence was attached - {imgstring}")
    else:
        writer("\nNot applicable here.")
        
    # --- Section 3: Containment ---
    writer("\n---\n## 3. Containment")
    containment_table_has_data = has_table_data(data, 'containment_table')
    strategy_has_data = not is_empty_value(data.get('containment_strategy'))
    if containment_table_has_data or strategy_has_data:
        if containment_table_has_data:
            write_table("containment_table", ["Date/Time (UTC)", "Action Taken", "Performed By", "System/Asset", "Outcome/Result"])
        if strategy_has_data:
            writer(f"\n**Containment Strategy:**\n{data['containment_strategy']}")
    else:
        writer("\nNot applicable here.")

    # --- Section 4: Eradication ---
    writer("\n---\n## 4. Eradication")
    eradication_table_has_data = has_table_data(data, 'eradication_table')
    eradication_summary_has_data = not is_empty_value(data.get('eradication_summary'))
    if eradication_table_has_data or eradication_summary_has_data:
        if eradication_table_has_data:
            write_table("eradication_table", ["Date/Time (UTC)", "Action Taken", "Performed By", "System/Asset", "Outcome/Result"])
        if eradication_summary_has_data:
            writer(f"\n**Eradication Summary:**\n{data['eradication_summary']}")
    else:
        writer("\nNot applicable here.")

    # --- Section 5: Recovery ---
    writer("\n---\n## 5. Recovery")
    recovery_table_has_data = has_table_data(data, 'recovery_table')
    recovery_plan_has_data = not is_empty_value(data.get('recovery_plan'))
    if recovery_table_has_data or recovery_plan_has_data:
        if recovery_table_has_data:
            write_table("recovery_table", ["Date/Time (UTC)", "Action Taken", "Performed By", "System/Asset", "Validation Steps"])
        if recovery_plan_has_data:
            writer(f"\n**Recovery Plan & Validation:**\n{data['recovery_plan']}")
    else:
        writer("\nNot applicable here.")

    # --- Section 6: Post-Incident ---
    writer("\n---\n## 6. Post-Incident Activities")
    timeline_has_data = not is_empty_value(data.get('timeline_summary'))
    root_cause_has_data = not is_empty_value(data.get('root_cause'))
    lessons_has_data = not is_empty_value(data.get('lessons_learned'))
    actions_table_has_data = has_table_data(data, 'actions_table')
    if timeline_has_data or root_cause_has_data or lessons_has_data or actions_table_has_data:
        if timeline_has_data:
            writer(f"**Incident Timeline Summary:**\n{data['timeline_summary']}")
        if root_cause_has_data:
            writer(f"\n**Root Cause Analysis:**\n{data['root_cause']}")
        if lessons_has_data:
            writer(f"\n**Lessons Learned:**\n{data['lessons_learned']}")
        if actions_table_has_data:
            writer("\n**Follow-up Actions:**")
            write_table("actions_table", ["Action Item", "Assigned To", "Due Date", "Status"])
    else:
        writer("\nNot applicable here.")

    # --- Section 7: Conclusion ---
    writer("\n---\n## 7. Conclusion")
    writer(f"**Incident Classification:** {data.get('conclusion_classification', 'N/A')}")
    writer(f"\n**Conclusion Summary:**\n{data.get('conclusion_summary', 'N/A')}")
    
    output.seek(0)
    return output

def generate_docx(data, files):
    """Generates a .docx file, showing all sections."""
    document = Document()
    document.add_heading(f"Incident Report: {data.get('incident_id', 'N/A')}", 0)

    def add_section(title, level=1):
        document.add_heading(title, level=level)

    def add_para(text, bold=False, inline_text=None):
        p = document.add_paragraph()
        p.add_run(text).bold = bold
        if inline_text:
            p.add_run(inline_text)

    # --- Helper to add tables ---
    def add_table(table_id, headers):
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        i = 0
        while True:
            row_key_base = f"{table_id}_{i}"
            if f"{row_key_base}_0" not in data: break
            row_data = [data.get(f"{row_key_base}_{j}", "").replace('T', ' ') for j in range(len(headers))]
            if any(cell.strip() for cell in row_data):
                row_cells = table.add_row().cells
                for j, cell_text in enumerate(row_data):
                    row_cells[j].text = cell_text
            i += 1
            
    # --- Section 1: Triage ---
    add_section("1. Triage & Initial Assessment")
    triage_fields = ['incident_id', 'reported_by', 'date_reported', 'priority', 'summary']
    triage_has_data = any(not is_empty_value(data.get(field)) for field in triage_fields)
    if triage_has_data:
        if not is_empty_value(data.get('incident_id')):
            add_para("Incident ID: ", bold=True, inline_text=data['incident_id'])
        if not is_empty_value(data.get('reported_by')):
            add_para("Reported By: ", bold=True, inline_text=data['reported_by'])
        if not is_empty_value(data.get('date_reported')):
            add_para("Date/Time Reported: ", bold=True, inline_text=f"{data['date_reported'].replace('T', ' ')} UTC")
        if not is_empty_value(data.get('priority')):
            add_para("Priority Level: ", bold=True, inline_text=data['priority'])
        if not is_empty_value(data.get('summary')):
            add_para("Initial Summary:", bold=True)
            document.add_paragraph(data['summary'])
    else:
        document.add_paragraph("Not applicable here.")

    # --- Section 2: Identification ---
    add_section("2. Identification & Scoping")
    id_table_has_data = has_table_data(data, 'identification_table')
    scope_has_data = not is_empty_value(data.get('scope_summary'))
    iocs_has_data = not is_empty_value(data.get('iocs'))
    if id_table_has_data or scope_has_data or iocs_has_data:
        if id_table_has_data:
            add_table("identification_table", ["Date/Time (UTC)", "Host/Asset", "Source IP", "Destination IP", "Event/Observation", "Analyst Notes"])
        if scope_has_data:
            add_para("\nSummary of Scope:", bold=True)
            document.add_paragraph(data['scope_summary'])
        if iocs_has_data:
            add_para("Indicators of Compromise (IoCs):", bold=True)
            document.add_paragraph(data['iocs'])
        if files and files[0].filename:
            add_section("Attached Evidence", level=2)
            for file in files:
                if file.filename:
                    file.seek(0)
                    document.add_paragraph(f"Filename: {file.filename}", style='Caption')
                    try:
                        document.add_picture(file, width=Inches(6.0))
                    except Exception as e:
                        document.add_paragraph(f"Could not embed image: {file.filename}. Error: {e}")
    else:
        document.add_paragraph("Not applicable here.")

    # --- Section 3: Containment ---
    add_section("3. Containment")
    containment_table_has_data = has_table_data(data, 'containment_table')
    strategy_has_data = not is_empty_value(data.get('containment_strategy'))
    if containment_table_has_data or strategy_has_data:
        if containment_table_has_data:
            add_table("containment_table", ["Date/Time (UTC)", "Action Taken", "Performed By", "System/Asset", "Outcome/Result"])
        if strategy_has_data:
            add_para("\nContainment Strategy:", bold=True)
            document.add_paragraph(data['containment_strategy'])
    else:
        document.add_paragraph("Not applicable here.")

    # --- Section 4: Eradication ---
    add_section("4. Eradication")
    eradication_table_has_data = has_table_data(data, 'eradication_table')
    eradication_summary_has_data = not is_empty_value(data.get('eradication_summary'))
    if eradication_table_has_data or eradication_summary_has_data:
        if eradication_table_has_data:
            add_table("eradication_table", ["Date/Time (UTC)", "Action Taken", "Performed By", "System/Asset", "Outcome/Result"])
        if eradication_summary_has_data:
            add_para("\nEradication Summary:", bold=True)
            document.add_paragraph(data['eradication_summary'])
    else:
        document.add_paragraph("Not applicable here.")

    # --- Section 5: Recovery ---
    add_section("5. Recovery")
    recovery_table_has_data = has_table_data(data, 'recovery_table')
    recovery_plan_has_data = not is_empty_value(data.get('recovery_plan'))
    if recovery_table_has_data or recovery_plan_has_data:
        if recovery_table_has_data:
            add_table("recovery_table", ["Date/Time (UTC)", "Action Taken", "Performed By", "System/Asset", "Validation Steps"])
        if recovery_plan_has_data:
            add_para("\nRecovery Plan & Validation:", bold=True)
            document.add_paragraph(data['recovery_plan'])
    else:
        document.add_paragraph("Not applicable here.")

    # --- Section 6: Post-Incident ---
    add_section("6. Post-Incident Activities")
    timeline_has_data = not is_empty_value(data.get('timeline_summary'))
    root_cause_has_data = not is_empty_value(data.get('root_cause'))
    lessons_has_data = not is_empty_value(data.get('lessons_learned'))
    actions_table_has_data = has_table_data(data, 'actions_table')
    if timeline_has_data or root_cause_has_data or lessons_has_data or actions_table_has_data:
        if timeline_has_data:
            add_para("Incident Timeline Summary:", bold=True)
            document.add_paragraph(data['timeline_summary'])
        if root_cause_has_data:
            add_para("Root Cause Analysis:", bold=True)
            document.add_paragraph(data['root_cause'])
        if lessons_has_data:
            add_para("Lessons Learned:", bold=True)
            document.add_paragraph(data['lessons_learned'])
        if actions_table_has_data:
            add_para("Follow-up Actions:", bold=True)
            add_table("actions_table", ["Action Item", "Assigned To", "Due Date", "Status"])
    else:
        document.add_paragraph("Not applicable here.")

    # --- Section 7: Conclusion ---
    add_section("7. Conclusion")
    add_para("Incident Classification: ", bold=True, inline_text=data.get('conclusion_classification', 'N/A'))
    add_para("Conclusion Summary:", bold=True)
    document.add_paragraph(data.get('conclusion_summary', 'N/A'))
    
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- FLASK ROUTES ---

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        form_data = request.form.to_dict()
        
        # Validate conclusion section before processing
        classification = form_data.get('conclusion_classification', '').strip()
        summary = form_data.get('conclusion_summary', '').strip()
        
        if not classification or not summary:
            # Return to form with error message
            now = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M")
            default_data = {
                "incident_id": form_data.get('incident_id', f"INC-{uuid.uuid4().hex[:8].upper()}"),
                "reported_by": form_data.get('reported_by', ''),
                "date_reported": form_data.get('date_reported', now),
                "priority": form_data.get('priority', 'Medium'),
                "summary": form_data.get('summary', ''),
                "validation_error": "Please complete the Conclusion section. Both Incident Classification and Conclusion Summary are required."
            }
            return render_template('index.html', data=default_data)
        
        files = request.files.getlist("screenshots")
        action = request.form.get('action')
        
        incident_id = form_data.get('incident_id', 'incident').strip()
        safe_filename = "".join([c for c in incident_id if c.isalpha() or c.isdigit() or c in ('_','-')]).rstrip()
        if not safe_filename: 
            safe_filename = "incident_report"

        if action == 'download_md':
            markdown_file = generate_markdown(form_data, files)
            return send_file(
                markdown_file,
                as_attachment=True,
                download_name=f'{safe_filename}.md',
                mimetype='text/markdown'
            )
        elif action == 'download_docx':
            docx_file = generate_docx(form_data, files)
            return send_file(
                docx_file,
                as_attachment=True,
                download_name=f'{safe_filename}.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    now = datetime.datetime.utcnow().strftime("%Y-%m-%dT%H:%M")
    default_data = {
        "incident_id": f"INC-{uuid.uuid4().hex[:8].upper()}",
        "reported_by": "",
        "date_reported": now,
        "priority": "Medium",
        "summary": ""
    }
    return render_template('index.html', data=default_data)

@app.route('/save_form', methods=['POST'])
def save_form():
    try:
        data = request.get_json()
        form_name = data.get('name')
        form_data = data.get('data')
        
        if not form_name or not form_data:
            return jsonify({'success': False, 'error': 'Missing form name or data'})
        
        # Create unique filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        form_id = f"{timestamp}_{uuid.uuid4().hex[:8]}"
        filename = f"{form_id}.json"
        filepath = os.path.join(SAVED_FORMS_DIR, filename)
        
        # Save form data
        save_data = {
            'id': form_id,
            'name': form_name,
            'timestamp': timestamp,
            'date': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'data': form_data
        }
        
        with open(filepath, 'w') as f:
            json.dump(save_data, f, indent=2)
        
        return jsonify({'success': True, 'form_id': form_id})
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/get_saved_forms', methods=['GET'])
def get_saved_forms():
    try:
        forms = []
        
        if os.path.exists(SAVED_FORMS_DIR):
            for filename in os.listdir(SAVED_FORMS_DIR):
                if filename.endswith('.json'):
                    filepath = os.path.join(SAVED_FORMS_DIR, filename)
                    try:
                        with open(filepath, 'r') as f:
                            form_data = json.load(f)
                            forms.append({
                                'id': form_data.get('id'),
                                'name': form_data.get('name'),
                                'date': form_data.get('date')
                            })
                    except Exception as e:
                        print(f"Error reading form {filename}: {e}")
                        continue
        
        # Sort by date (newest first)
        forms.sort(key=lambda x: x.get('date', ''), reverse=True)
        
        return jsonify({'success': True, 'forms': forms})
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/load_form/<form_id>', methods=['GET'])
def load_form(form_id):
    try:
        # Find the form file
        form_file = None
        if os.path.exists(SAVED_FORMS_DIR):
            for filename in os.listdir(SAVED_FORMS_DIR):
                if filename.startswith(form_id) and filename.endswith('.json'):
                    form_file = os.path.join(SAVED_FORMS_DIR, filename)
                    break
        
        if not form_file or not os.path.exists(form_file):
            return jsonify({'success': False, 'error': 'Form not found'})
        
        with open(form_file, 'r') as f:
            form_data = json.load(f)
        
        return jsonify({'success': True, 'data': form_data.get('data', {})})
    
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

if __name__ == '__main__':
    # To run this:
    # 1. Make sure you have Flask and python-docx installed:
    #    pip install Flask python-docx
    # 2. Save the code as a Python file (e.g., app.py).
    # 3. Run from your terminal: python app.py
    # 4. Open your web browser to http://127.0.0.1:5000
    app.run(debug=False)